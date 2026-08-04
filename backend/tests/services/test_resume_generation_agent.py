"""
Tests for ResumeGenerationAgent.

The agent is a controller, not a single LLM call:
- match_keywords and compute_ats run deterministically (no model call).
- Exactly one Gemini call optimizes the resume content — it returns
  ContentPatch[] keyed by block_id (see synthetic_profile_layout.py), never a
  restructured resume object directly. A deterministic post-step (Patch
  Engine + flattening) reconstructs the OptimizedResume-shaped dict.
- The returned dict matches the documented pipeline output shape.
"""
import io
import json
from unittest.mock import AsyncMock, MagicMock, patch

import docx
import pytest

from services.docx_layout_extractor import extract_docx_layout
from services.llm_output import LLMOutputError
from services.resume_generation_agent import _SYSTEM_PROMPT, ResumeGenerationAgent


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_response(data: dict) -> MagicMock:
    mock = MagicMock()
    mock.text = json.dumps(data)
    return mock


_PROFILE = {
    "headline": "Backend Engineer",
    "technical_skills": ["Python", "AWS"],
    "work_experience": [
        {
            "title": "Software Engineer",
            "company": "Acme Corp",
            "start_date": "2020",
            "end_date": "Present",
            "bullets": ["Built APIs with Python and AWS"],
            "technologies": ["Python", "AWS"],
        }
    ],
    "projects": [],
}

_JOB = {"skills": ["Python", "AWS", "Terraform"]}

_PATCHES_RESPONSE = {
    "patches": [
        {"block_id": "headline", "new_text": "Senior Backend Engineer"},
        {"block_id": "summary", "new_text": "Backend engineer with Python and AWS experience."},
        {"block_id": "skills", "new_text": "Python, AWS"},
        {"block_id": "work_experience[0].bullets[0]", "new_text": "Built scalable APIs with Python and AWS"},
        {
            "block_id": "changes_summary",
            "new_text": "Emphasized your Python and AWS experience since the job requires both.\n"
                         "Could not address Terraform — no matching experience was found in your profile.",
        },
    ]
}

_EXPECTED_OPTIMIZED_RESUME = {
    "name": None,
    "email": None,
    "phone": None,
    "headline": "Senior Backend Engineer",
    "summary": "Backend engineer with Python and AWS experience.",
    "links": {},
    "skills": ["Python", "AWS"],
    "skill_groups": [{"label": "Additional", "items": ["Python", "AWS"]}],
    "experience": [
        {
            "title": "Software Engineer",
            "company": "Acme Corp",
            "location": None,
            "start_date": "2020",
            "end_date": "Present",
            "bullets": ["Built scalable APIs with Python and AWS"],
        }
    ],
    "projects": [],
    "education": [],
    "certifications": [],
    "awards": [],
    "leadership": [],
    "volunteering": [],
    "publications": [],
    "interests": [],
    "references": [],
    "changes_summary": [
        "Emphasized your Python and AWS experience since the job requires both.",
        "Could not address Terraform — no matching experience was found in your profile.",
    ],
}


@pytest.fixture
def mock_genai():
    """Patches genai in the agent module and yields the mock client."""
    with patch("services.resume_generation_agent.genai") as patched:
        mock_client = MagicMock()
        patched.Client.return_value = mock_client
        mock_client.aio.models.generate_content = AsyncMock()
        yield mock_client


# ---------------------------------------------------------------------------
# Pipeline shape
# ---------------------------------------------------------------------------

@pytest.mark.anyio
async def test_returns_full_pipeline_output_shape(mock_genai):
    mock_genai.aio.models.generate_content.return_value = _make_response(_PATCHES_RESPONSE)

    result = await ResumeGenerationAgent().generate(_PROFILE, _JOB)

    assert set(result.keys()) == {
        "ats_score", "matched_keywords", "missing_keywords", "added_keywords", "optimized_resume",
        "patches", "render_layout", "layout_preserved", "report",
    }
    assert result["optimized_resume"] == _EXPECTED_OPTIMIZED_RESUME
    assert result["patches"] == _PATCHES_RESPONSE["patches"]
    assert result["render_layout"] is None  # no layout_document was given
    assert result["layout_preserved"] is False
    assert set(result["report"].keys()) == {
        "ats_score_before", "ats_score_after", "matched_keywords_before", "matched_keywords_after",
        "keywords_added", "keywords_skipped", "unused_candidate_skills", "skills_reordered",
        "summary_rewritten", "experience_bullets_modified", "projects_modified", "highlights",
        "rewrite_similarity_avg", "rewrite_quality_issues", "project_ranking",
    }


@pytest.mark.anyio
async def test_computes_ats_score_deterministically(mock_genai):
    mock_genai.aio.models.generate_content.return_value = _make_response(_PATCHES_RESPONSE)

    result = await ResumeGenerationAgent().generate(_PROFILE, _JOB)

    # 2 of 3 job keywords (Python, AWS) are present in the profile.
    assert result["ats_score"] == pytest.approx(66.67)
    assert result["matched_keywords"] == ["Python", "AWS"]
    assert result["missing_keywords"] == ["Terraform"]


@pytest.mark.anyio
async def test_added_keywords_excludes_ones_with_no_basis_in_the_profile(mock_genai):
    """A missing_keyword the LLM wove into wording only counts as genuinely
    added if the candidate profile has real evidence for it somewhere (see
    keyword_matcher.filter_backed_keywords) — otherwise it's a fabrication
    the prompt's "never invent" rule should have caught but didn't."""
    profile = {
        "technical_skills": ["Python"],
        "work_experience": [
            {
                "title": "Software Engineer", "company": "Acme Corp",
                "bullets": ["Deployed services with Kubernetes"], "technologies": ["Python"],
            }
        ],
        "projects": [],
    }
    job = {"skills": ["Python", "Kubernetes", "Terraform"]}
    patches = {
        "patches": [
            {"block_id": "headline", "new_text": ""},
            {"block_id": "summary", "new_text": "Backend engineer experienced with Kubernetes and Terraform."},
            {"block_id": "skills", "new_text": "Python"},
            {"block_id": "work_experience[0].bullets[0]", "new_text": "Deployed services with Kubernetes"},
            {"block_id": "changes_summary", "new_text": "Wove in Kubernetes and Terraform."},
        ]
    }
    mock_genai.aio.models.generate_content.return_value = _make_response(patches)

    result = await ResumeGenerationAgent().generate(profile, job)

    assert result["added_keywords"] == ["Kubernetes"]


@pytest.mark.anyio
async def test_calls_model_exactly_once(mock_genai):
    mock_genai.aio.models.generate_content.return_value = _make_response(_PATCHES_RESPONSE)

    await ResumeGenerationAgent().generate(_PROFILE, _JOB)

    assert mock_genai.aio.models.generate_content.call_count == 1


# ---------------------------------------------------------------------------
# Optimization report
# ---------------------------------------------------------------------------

@pytest.mark.anyio
async def test_report_computes_before_after_scores_and_pass_through_reasoning(mock_genai):
    profile = {
        "technical_skills": ["Python"],
        "devops_tools": ["Docker"],  # not relevant to this job — should surface as unused
        "work_experience": [
            {
                "title": "Software Engineer", "company": "Acme Corp",
                "bullets": ["Built APIs with Python"], "technologies": ["Python"],
            }
        ],
        "projects": [],
    }
    job = {"skills": ["Python", "Kubernetes", "Terraform"]}
    mock_genai.aio.models.generate_content.return_value = _make_response({
        "patches": [
            {"block_id": "headline", "new_text": ""},
            {"block_id": "summary", "new_text": "Backend engineer experienced with Python."},
            {"block_id": "skills", "new_text": "Python"},
            {"block_id": "work_experience[0].bullets[0]", "new_text": "Built scalable APIs with Python"},
        ],
        "highlights": [
            {"section": "Experience", "summary": "Reworded bullet for clarity.", "impact": "medium"},
        ],
        "keywords_skipped": [
            {"keyword": "Terraform", "reason": "No demonstrated Terraform experience in the profile."},
        ],
    })

    result = await ResumeGenerationAgent().generate(profile, job)
    report = result["report"]

    # 1 of 3 job keywords (Python) is present; neither Kubernetes nor
    # Terraform got woven into the wording, so before/after are identical.
    assert report["ats_score_before"] == pytest.approx(33.33)
    assert report["ats_score_after"] == pytest.approx(33.33)
    assert report["matched_keywords_before"] == 1
    assert report["matched_keywords_after"] == 1
    assert report["keywords_added"] == []
    assert report["keywords_skipped"] == [
        {"keyword": "Terraform", "reason": "No demonstrated Terraform experience in the profile."},
    ]
    assert report["unused_candidate_skills"] == ["Docker"]
    assert report["skills_reordered"] is False
    assert report["summary_rewritten"] is True
    assert report["experience_bullets_modified"] == 1
    assert report["projects_modified"] == 0
    assert report["highlights"] == [
        {"section": "Experience", "summary": "Reworded bullet for clarity.", "impact": "medium"},
    ]


@pytest.mark.anyio
async def test_report_after_score_credits_a_genuinely_woven_in_keyword(mock_genai):
    profile = {
        "technical_skills": ["Python"],
        "work_experience": [
            {
                "title": "Software Engineer", "company": "Acme Corp",
                "bullets": ["Deployed services with Kubernetes"], "technologies": ["Python"],
            }
        ],
        "projects": [],
    }
    job = {"skills": ["Python", "Kubernetes"]}
    mock_genai.aio.models.generate_content.return_value = _make_response({
        "patches": [
            {"block_id": "headline", "new_text": ""},
            {"block_id": "summary", "new_text": ""},
            {"block_id": "skills", "new_text": "Python"},
            {"block_id": "work_experience[0].bullets[0]", "new_text": "Deployed services with Kubernetes"},
        ],
    })

    result = await ResumeGenerationAgent().generate(profile, job)
    report = result["report"]

    assert report["ats_score_before"] == pytest.approx(50.0)  # 1 of 2
    assert report["ats_score_after"] == pytest.approx(100.0)  # Kubernetes now credited
    assert report["matched_keywords_after"] == 2
    assert report["keywords_added"] == ["Kubernetes"]


@pytest.mark.anyio
async def test_report_detects_skills_reordering(mock_genai):
    profile = {
        "technical_skills": ["Python", "Go"],
        "programming_languages": ["Java"],
        "work_experience": [],
        "projects": [],
    }
    job = {"skills": ["Java"]}
    mock_genai.aio.models.generate_content.return_value = _make_response({
        "patches": [
            {"block_id": "headline", "new_text": ""},
            {"block_id": "summary", "new_text": ""},
            # Reordered to lead with the job-relevant skill, unlike the
            # profile's own field order (technical_skills before
            # programming_languages).
            {"block_id": "skills", "new_text": "Java, Python, Go"},
        ],
    })

    result = await ResumeGenerationAgent().generate(profile, job)

    assert result["report"]["skills_reordered"] is True


@pytest.mark.anyio
async def test_report_includes_project_ranking_reasons(mock_genai):
    profile = {
        "technical_skills": ["Python"],
        "work_experience": [],
        "projects": [
            {"name": "Relevant Project", "technologies": ["Azure", "Docker"]},
            {"name": "Irrelevant Project", "technologies": ["COBOL"]},
        ],
    }
    job = {"skills": ["Azure", "Docker"]}
    mock_genai.aio.models.generate_content.return_value = _make_response({
        "patches": [
            {"block_id": "headline", "new_text": ""},
            {"block_id": "summary", "new_text": ""},
            {"block_id": "skills", "new_text": "Python"},
            {"block_id": "projects[0].description", "new_text": ""},
            {"block_id": "projects[0].technologies", "new_text": "Azure, Docker"},
            {"block_id": "projects[1].description", "new_text": ""},
            {"block_id": "projects[1].technologies", "new_text": "COBOL"},
        ],
    })

    result = await ResumeGenerationAgent().generate(profile, job)

    assert result["report"]["project_ranking"] == [
        {"name": "Relevant Project", "matched_on": ["Azure", "Docker"]},
        {"name": "Irrelevant Project", "matched_on": []},
    ]


@pytest.mark.anyio
async def test_report_counts_modified_project_entries(mock_genai):
    profile = {
        "technical_skills": ["Python"],
        "work_experience": [],
        "projects": [
            {"name": "Project A", "description": "Built a tool.", "technologies": ["Python"]},
            {"name": "Project B", "description": "Built another tool.", "technologies": ["Python"]},
        ],
    }
    job = {"skills": ["Python"]}
    mock_genai.aio.models.generate_content.return_value = _make_response({
        "patches": [
            {"block_id": "headline", "new_text": ""},
            {"block_id": "summary", "new_text": ""},
            {"block_id": "skills", "new_text": "Python"},
            {"block_id": "projects[0].description", "new_text": "Built a genuinely useful tool."},
            {"block_id": "projects[0].technologies", "new_text": "Python"},
            {"block_id": "projects[1].description", "new_text": "Built another tool."},  # unchanged
            {"block_id": "projects[1].technologies", "new_text": "Python"},
        ],
    })

    result = await ResumeGenerationAgent().generate(profile, job)

    assert result["report"]["projects_modified"] == 1


# ---------------------------------------------------------------------------
# Inferred keywords (see inferable_keywords.py / _augment_with_inferred_keywords)
# ---------------------------------------------------------------------------

@pytest.mark.anyio
async def test_missing_keywords_includes_an_inferred_keyword_the_job_cares_about(mock_genai):
    """The job lists "Agile" as a required skill; nothing in the candidate's
    structured technologies/skills_demonstrated fields says "agile", but a
    real bullet describes sprint-based work — that's real evidence, so the
    keyword should be offered to the optimizer as missing, not silently
    invisible the way it would be under literal tag-only matching."""
    profile = {
        "technical_skills": ["Python"],
        "work_experience": [
            {
                "title": "Engineer", "company": "Acme",
                "bullets": ["Delivered features in two-week agile sprints using Python"],
                "technologies": ["Python"],
            }
        ],
        "projects": [],
    }
    job = {"skills": ["Python", "Agile"]}
    mock_genai.aio.models.generate_content.return_value = _make_response(_PATCHES_RESPONSE)

    result = await ResumeGenerationAgent().generate(profile, job)

    assert "Agile Software Development" in result["missing_keywords"]


@pytest.mark.anyio
async def test_inferred_keyword_omitted_when_job_does_not_care(mock_genai):
    """Same real evidence (agile sprints in the bullet text), but this job
    never mentions Agile/Scrum/sprints at all — nothing should be offered
    that the job itself doesn't ask for."""
    profile = {
        "technical_skills": ["Python"],
        "work_experience": [
            {
                "title": "Engineer", "company": "Acme",
                "bullets": ["Delivered features in two-week agile sprints using Python"],
                "technologies": ["Python"],
            }
        ],
        "projects": [],
    }
    job = {"skills": ["Python"]}
    mock_genai.aio.models.generate_content.return_value = _make_response(_PATCHES_RESPONSE)

    result = await ResumeGenerationAgent().generate(profile, job)

    assert "Agile Software Development" not in result["missing_keywords"]


@pytest.mark.anyio
async def test_inferred_keyword_omitted_when_no_evidence_exists(mock_genai):
    """The job wants Agile experience, but nothing in the candidate's real
    text supports it — never offer an unbacked keyword, even a soft/process
    one, as something the optimizer might weave in."""
    profile = {
        "technical_skills": ["Python"],
        "work_experience": [
            {
                "title": "Engineer", "company": "Acme",
                "bullets": ["Wrote a script to rename files"],
                "technologies": ["Python"],
            }
        ],
        "projects": [],
    }
    job = {"skills": ["Python", "Agile"]}
    mock_genai.aio.models.generate_content.return_value = _make_response(_PATCHES_RESPONSE)

    result = await ResumeGenerationAgent().generate(profile, job)

    assert "Agile Software Development" not in result["missing_keywords"]


# ---------------------------------------------------------------------------
# Rewrite quality measurement (see _rewrite_quality)
# ---------------------------------------------------------------------------

@pytest.mark.anyio
async def test_rewrite_similarity_avg_is_none_when_nothing_changed(mock_genai):
    profile = {
        "technical_skills": ["Python"],
        "work_experience": [
            {"title": "Engineer", "company": "Acme", "bullets": ["Built APIs with Python"], "technologies": ["Python"]}
        ],
        "projects": [],
    }
    job = {"skills": ["Python"]}
    mock_genai.aio.models.generate_content.return_value = _make_response({
        "patches": [
            {"block_id": "headline", "new_text": ""},
            {"block_id": "summary", "new_text": ""},
            {"block_id": "skills", "new_text": "Python"},
            # Bullet comes back byte-identical — nothing actually changed.
            {"block_id": "work_experience[0].bullets[0]", "new_text": "Built APIs with Python"},
        ],
    })

    result = await ResumeGenerationAgent().generate(profile, job)

    assert result["report"]["rewrite_similarity_avg"] is None
    assert result["report"]["rewrite_quality_issues"] == []


@pytest.mark.anyio
async def test_rewrite_similarity_avg_reflects_a_real_rewrite(mock_genai):
    profile = {
        "technical_skills": ["Python"],
        "work_experience": [
            {"title": "Engineer", "company": "Acme", "bullets": ["Built APIs with Python"], "technologies": ["Python"]}
        ],
        "projects": [],
    }
    job = {"skills": ["Python"]}
    mock_genai.aio.models.generate_content.return_value = _make_response({
        "patches": [
            {"block_id": "headline", "new_text": ""},
            {"block_id": "summary", "new_text": ""},
            {"block_id": "skills", "new_text": "Python"},
            {"block_id": "work_experience[0].bullets[0]",
             "new_text": "Designed and shipped production REST APIs serving internal teams"},
        ],
    })

    result = await ResumeGenerationAgent().generate(profile, job)

    assert result["report"]["rewrite_similarity_avg"] is not None
    assert result["report"]["rewrite_similarity_avg"] < 0.9


@pytest.mark.anyio
async def test_rewrite_quality_flags_a_keyword_stitched_on_with_minimal_change(mock_genai):
    """A block that claims to now carry an added keyword but barely changed
    in wording is a signal the keyword was appended rather than woven in —
    see the STEP 1/STEP 2 editing philosophy in _SYSTEM_PROMPT."""
    profile = {
        "technical_skills": ["Python"],
        "work_experience": [
            {
                "title": "Engineer", "company": "Acme",
                # "Kubernetes" is only in the bullet's own prose, never tagged
                # in technologies — so it starts out as a missing_keyword
                # even though it's already truthfully present in the text.
                "bullets": ["Built APIs with Python and deployed services with Kubernetes"],
                "technologies": ["Python"],
            }
        ],
        "projects": [],
    }
    job = {"skills": ["Python", "Kubernetes"]}
    mock_genai.aio.models.generate_content.return_value = _make_response({
        "patches": [
            {"block_id": "headline", "new_text": ""},
            {"block_id": "summary", "new_text": ""},
            {"block_id": "skills", "new_text": "Python"},
            # Near-identical to the original — just a trailing period swap —
            # while still "carrying" Kubernetes (it was already in the text).
            {"block_id": "work_experience[0].bullets[0]",
             "new_text": "Built APIs with Python and deployed services with Kubernetes,"},
        ],
    })

    result = await ResumeGenerationAgent().generate(profile, job)

    assert any(
        "work_experience[0].bullets[0]" in issue for issue in result["report"]["rewrite_quality_issues"]
    )


def test_prompt_defines_a_two_step_rewrite_decision_framework():
    """Regression guard: without an explicit, quantified decision rule, the
    model tends to default to minimal cosmetic edits regardless of genuine
    keyword gaps (see the review that motivated this framework) — a future
    prompt edit must not silently drop it."""
    normalized_prompt = " ".join(_SYSTEM_PROMPT.split())
    assert "STEP 1" in normalized_prompt
    assert "STEP 1's \"yes\" branch" in normalized_prompt or "STEP 1" in normalized_prompt
    assert "light polish only" in normalized_prompt


def test_prompt_reconciles_style_preservation_with_the_rewrite_framework():
    """Regression guard: priority item 4 ("preserve writing style") predates
    the STEP 1/STEP 2 framework and, left unqualified, silently conflicts
    with it — the model could use "the sentence already works" to justify
    not rewriting a block STEP 2 says genuinely needs it. A future prompt
    edit must not drop the explicit reconciliation."""
    normalized_prompt = " ".join(_SYSTEM_PROMPT.split())
    assert "Preserve writing style ONLY when STEP 1 above found no applicable" in normalized_prompt


def test_prompt_instructs_surfacing_unused_quantified_detail():
    """Regression guard: a bullet can understate its own impact even when a
    real number for the same fact exists elsewhere in candidate_profile —
    a future prompt edit must not silently drop the instruction to use it
    (never invent one)."""
    normalized_prompt = " ".join(_SYSTEM_PROMPT.split())
    assert "understates or omits" in normalized_prompt
    assert "Never invent a number or" in normalized_prompt


def test_prompt_instructs_differentiating_thematically_similar_projects():
    """Regression guard: without this, two genuinely distinct projects that
    happen to share a high-level pitch (e.g. two stock-valuation tools) can
    both get rewritten around the same angle, reading as repetition on the
    resume — a future prompt edit must not silently drop this guidance."""
    normalized_prompt = " ".join(_SYSTEM_PROMPT.split())
    assert "differentiate them by each project's actual" in normalized_prompt


def test_prompt_summary_rule_names_generic_filler_to_avoid_when_retargeting():
    """Regression guard: retargeting a summary for a specific job must not
    be an excuse to flatten it back into generic phrasing — a future prompt
    edit must not silently drop the explicit filler list here, matching
    candidate_profile_agent.py's own list."""
    normalized_prompt = " ".join(_SYSTEM_PROMPT.split())
    assert "eager to contribute to a dynamic team" in normalized_prompt


def test_prompt_length_target_does_not_override_step_two_rewrites():
    """Regression guard: the experience-bullet length guidance used to say
    "never noticeably longer than the bullet's current length" with no
    exception — read literally, that silently overrides STEP 2's
    restructuring for any block that genuinely needs a few more words. A
    future prompt edit must not reintroduce a hard ceiling here."""
    normalized_prompt = " ".join(_SYSTEM_PROMPT.split())
    assert "a length target, not a hard ceiling" in normalized_prompt


# ---------------------------------------------------------------------------
# Resilience — patches referencing unknown/missing block_ids
# ---------------------------------------------------------------------------

@pytest.mark.anyio
async def test_unknown_block_id_in_response_is_ignored_not_raised(mock_genai):
    response = {"patches": [*_PATCHES_RESPONSE["patches"], {"block_id": "not_a_real_block", "new_text": "ghost"}]}
    mock_genai.aio.models.generate_content.return_value = _make_response(response)

    result = await ResumeGenerationAgent().generate(_PROFILE, _JOB)

    assert result["optimized_resume"] == _EXPECTED_OPTIMIZED_RESUME


@pytest.mark.anyio
async def test_missing_patch_for_a_block_keeps_its_original_placeholder_text(mock_genai):
    # The model never patches "headline" at all.
    patches = [p for p in _PATCHES_RESPONSE["patches"] if p["block_id"] != "headline"]
    mock_genai.aio.models.generate_content.return_value = _make_response({"patches": patches})

    result = await ResumeGenerationAgent().generate(_PROFILE, _JOB)

    assert result["optimized_resume"]["headline"] == "Backend Engineer"  # original profile value, untouched


# ---------------------------------------------------------------------------
# Model input
# ---------------------------------------------------------------------------

@pytest.mark.anyio
async def test_sends_matched_and_missing_keywords_to_model(mock_genai):
    mock_genai.aio.models.generate_content.return_value = _make_response(_PATCHES_RESPONSE)

    await ResumeGenerationAgent().generate(_PROFILE, _JOB)

    contents = json.loads(mock_genai.aio.models.generate_content.call_args.kwargs["contents"])
    assert contents["matched_keywords"] == ["Python", "AWS"]
    assert contents["missing_keywords"] == ["Terraform"]
    assert contents["job"] == _JOB


@pytest.mark.anyio
async def test_sends_editable_blocks_with_block_ids_and_placeholder_text(mock_genai):
    mock_genai.aio.models.generate_content.return_value = _make_response(_PATCHES_RESPONSE)

    await ResumeGenerationAgent().generate(_PROFILE, _JOB)

    contents = json.loads(mock_genai.aio.models.generate_content.call_args.kwargs["contents"])
    editable_blocks = {block["block_id"]: block["text"] for block in contents["editable_blocks"]}
    assert editable_blocks["headline"] == "Backend Engineer"
    assert editable_blocks["work_experience[0].bullets[0]"] == "Built APIs with Python and AWS"
    assert editable_blocks["skills"] == ""  # blank canvas, synthesized fresh


@pytest.mark.anyio
async def test_sends_full_candidate_profile_as_context(mock_genai):
    """The model must still see the whole ranked profile (all skill
    categories, full entries) even though it can now only ever return
    wording changes through editable_blocks."""
    mock_genai.aio.models.generate_content.return_value = _make_response(_PATCHES_RESPONSE)

    await ResumeGenerationAgent().generate(_PROFILE, _JOB)

    contents = json.loads(mock_genai.aio.models.generate_content.call_args.kwargs["contents"])
    assert contents["candidate_profile"]["technical_skills"] == ["Python", "AWS"]


# ---------------------------------------------------------------------------
# layout_document — real in-place rendering via profile_layout_correlator.py
# ---------------------------------------------------------------------------

def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _make_docx_bytes_with_styles(paragraphs: list[tuple[str, str | None]]) -> bytes:
    document = docx.Document()
    for text, style_name in paragraphs:
        document.add_paragraph(text, style=style_name)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.anyio
async def test_confident_correlation_produces_a_patched_real_layout(mock_genai):
    # The real document's text closely matches the profile fields it was
    # extracted from, so correlation should confidently find every block.
    docx_bytes = _make_docx_bytes([
        "Backend Engineer",
        "Built APIs with Python and AWS",
    ])
    layout_document = extract_docx_layout(docx_bytes).model_dump()
    profile = {
        "headline": "Backend Engineer",
        "work_experience": [{"bullets": ["Built APIs with Python and AWS"]}],
    }
    mock_genai.aio.models.generate_content.return_value = _make_response({
        "patches": [
            {"block_id": "headline", "new_text": "Senior Backend Engineer"},
            {"block_id": "summary", "new_text": ""},
            {"block_id": "skills", "new_text": ""},
            {"block_id": "work_experience[0].bullets[0]", "new_text": "Built scalable APIs with Python and AWS"},
            {"block_id": "changes_summary", "new_text": ""},
        ]
    })

    result = await ResumeGenerationAgent().generate(profile, _JOB, layout_document=layout_document)

    assert result["layout_preserved"] is True
    real_blocks = {
        block["block_id"]: block["text"]
        for section in result["render_layout"]["sections"]
        for block in section["blocks"]
    }
    assert real_blocks["paragraph[0]"] == "Senior Backend Engineer"
    assert real_blocks["paragraph[1]"] == "Built scalable APIs with Python and AWS"


@pytest.mark.anyio
async def test_skills_patch_is_distributed_across_overflow_blocks(mock_genai):
    # A multi-line skills section: instead of writing the whole list to the
    # first line and blanking the rest, each line gets its own roughly-even
    # share so no single line has to fit what used to span several.
    docx_bytes = _make_docx_bytes_with_styles([
        ("Backend Engineer", None),
        ("Skills", "Heading 1"),
        ("Python", None),
        ("PostgreSQL", None),
        ("Built APIs with Python and AWS", None),
    ])
    layout_document = extract_docx_layout(docx_bytes).model_dump()
    profile = {
        "headline": "Backend Engineer",
        "work_experience": [{"bullets": ["Built APIs with Python and AWS"]}],
    }
    mock_genai.aio.models.generate_content.return_value = _make_response({
        "patches": [
            {"block_id": "headline", "new_text": "Senior Backend Engineer"},
            {"block_id": "skills", "new_text": "Python, PostgreSQL, Docker"},
            {"block_id": "work_experience[0].bullets[0]", "new_text": "Built scalable APIs with Python and AWS"},
        ]
    })

    result = await ResumeGenerationAgent().generate(profile, _JOB, layout_document=layout_document)

    real_blocks = {
        block["block_id"]: block["text"]
        for section in result["render_layout"]["sections"]
        for block in section["blocks"]
    }
    assert real_blocks["paragraph[2]"] == "Python, PostgreSQL"
    assert real_blocks["paragraph[3]"] == "Docker"


@pytest.mark.anyio
async def test_skills_patch_distributed_across_three_blocks_loses_no_items(mock_genai):
    # Mirrors a real "Technical Skills" section: heading + 3 category lines.
    # Every item from the LLM's compiled list must land in exactly one
    # block, with no item dropped or duplicated across the split.
    docx_bytes = _make_docx_bytes_with_styles([
        ("Backend Engineer", None),
        ("Skills", "Heading 1"),
        ("Python", None),
        ("PostgreSQL", None),
        ("Docker", None),
        ("Built APIs with Python and AWS", None),
    ])
    layout_document = extract_docx_layout(docx_bytes).model_dump()
    profile = {
        "headline": "Backend Engineer",
        "work_experience": [{"bullets": ["Built APIs with Python and AWS"]}],
    }
    skills_items = ["Python", "Java", "SQL", "HTML5", "CSS", "JavaScript", "TypeScript"]
    mock_genai.aio.models.generate_content.return_value = _make_response({
        "patches": [
            {"block_id": "headline", "new_text": "Senior Backend Engineer"},
            {"block_id": "skills", "new_text": ", ".join(skills_items)},
            {"block_id": "work_experience[0].bullets[0]", "new_text": "Built scalable APIs with Python and AWS"},
        ]
    })

    result = await ResumeGenerationAgent().generate(profile, _JOB, layout_document=layout_document)

    real_blocks = {
        block["block_id"]: block["text"]
        for section in result["render_layout"]["sections"]
        for block in section["blocks"]
    }
    skills_block_ids = ["paragraph[2]", "paragraph[3]", "paragraph[4]"]
    chunks = [real_blocks[block_id] for block_id in skills_block_ids]
    assert all(chunk for chunk in chunks)  # every block got a non-empty share
    recombined = [item.strip() for chunk in chunks for item in chunk.split(",")]
    assert recombined == skills_items


@pytest.mark.anyio
async def test_single_block_skills_section_gets_the_full_list_unsplit(mock_genai):
    docx_bytes = _make_docx_bytes_with_styles([
        ("Backend Engineer", None),
        ("Skills", "Heading 1"),
        ("Python", None),
        ("Built APIs with Python and AWS", None),
    ])
    layout_document = extract_docx_layout(docx_bytes).model_dump()
    profile = {
        "headline": "Backend Engineer",
        "work_experience": [{"bullets": ["Built APIs with Python and AWS"]}],
    }
    mock_genai.aio.models.generate_content.return_value = _make_response({
        "patches": [
            {"block_id": "headline", "new_text": "Senior Backend Engineer"},
            {"block_id": "skills", "new_text": "Python, PostgreSQL, Docker"},
            {"block_id": "work_experience[0].bullets[0]", "new_text": "Built scalable APIs with Python and AWS"},
        ]
    })

    result = await ResumeGenerationAgent().generate(profile, _JOB, layout_document=layout_document)

    real_blocks = {
        block["block_id"]: block["text"]
        for section in result["render_layout"]["sections"]
        for block in section["blocks"]
    }
    assert real_blocks["paragraph[2]"] == "Python, PostgreSQL, Docker"


@pytest.mark.anyio
async def test_skills_overflow_blocks_are_untouched_without_a_skills_patch(mock_genai):
    # Same multi-line skills section, but the model never emits a "skills"
    # patch — the overflow blocks must be left exactly as-is, not blanked.
    docx_bytes = _make_docx_bytes_with_styles([
        ("Backend Engineer", None),
        ("Skills", "Heading 1"),
        ("Python", None),
        ("PostgreSQL", None),
        ("Built APIs with Python and AWS", None),
    ])
    layout_document = extract_docx_layout(docx_bytes).model_dump()
    profile = {
        "headline": "Backend Engineer",
        "work_experience": [{"bullets": ["Built APIs with Python and AWS"]}],
    }
    mock_genai.aio.models.generate_content.return_value = _make_response({
        "patches": [
            {"block_id": "headline", "new_text": "Senior Backend Engineer"},
            {"block_id": "work_experience[0].bullets[0]", "new_text": "Built scalable APIs with Python and AWS"},
        ]
    })

    result = await ResumeGenerationAgent().generate(profile, _JOB, layout_document=layout_document)

    real_blocks = {
        block["block_id"]: block["text"]
        for section in result["render_layout"]["sections"]
        for block in section["blocks"]
    }
    assert real_blocks["paragraph[2]"] == "Python"
    assert real_blocks["paragraph[3]"] == "PostgreSQL"


@pytest.mark.anyio
async def test_low_confidence_correlation_falls_back_without_a_render_layout(mock_genai):
    # Nothing in this document resembles the profile's fields at all.
    docx_bytes = _make_docx_bytes(["Completely unrelated filler text."])
    layout_document = extract_docx_layout(docx_bytes).model_dump()
    profile = {
        "headline": "Backend Engineer",
        "work_experience": [{"bullets": ["Built APIs with Python and AWS"]}],
    }
    mock_genai.aio.models.generate_content.return_value = _make_response(_PATCHES_RESPONSE)

    result = await ResumeGenerationAgent().generate(profile, _JOB, layout_document=layout_document)

    assert result["layout_preserved"] is False
    assert result["render_layout"] is None


@pytest.mark.anyio
async def test_no_layout_document_means_no_render_layout(mock_genai):
    mock_genai.aio.models.generate_content.return_value = _make_response(_PATCHES_RESPONSE)

    result = await ResumeGenerationAgent().generate(_PROFILE, _JOB, layout_document=None)

    assert result["layout_preserved"] is False
    assert result["render_layout"] is None


# ---------------------------------------------------------------------------
# Model configuration
# ---------------------------------------------------------------------------

@pytest.mark.anyio
async def test_uses_flash_lite_model(mock_genai):
    mock_genai.aio.models.generate_content.return_value = _make_response(_PATCHES_RESPONSE)

    await ResumeGenerationAgent().generate(_PROFILE, _JOB)

    assert mock_genai.aio.models.generate_content.call_args.kwargs["model"] == "gemini-2.5-flash-lite"


@pytest.mark.anyio
async def test_uses_json_response_mode_and_zero_temperature(mock_genai):
    mock_genai.aio.models.generate_content.return_value = _make_response(_PATCHES_RESPONSE)

    await ResumeGenerationAgent().generate(_PROFILE, _JOB)

    config = mock_genai.aio.models.generate_content.call_args.kwargs["config"]
    assert config.response_mime_type == "application/json"
    assert config.temperature == 0


def test_prompt_tells_model_to_leave_changes_summary_blank():
    """Regression guard: changes_summary is replaced by the structured
    highlights/keywords_skipped report — the prompt must still mention the
    block by name (it's still a block_id the model is given) so the model
    knows to blank it rather than resume writing prose into it."""
    normalized_prompt = " ".join(_SYSTEM_PROMPT.split())
    assert "changes_summary, which this pipeline no longer displays" in normalized_prompt


def test_prompt_asks_for_grouped_highlights_and_keyword_skip_reasons():
    """Regression guard: this is what fixes the old per-block "I emphasized
    X" / "I highlighted Y" repetitive changes_summary output — a future
    prompt edit must not silently drop the structured replacement."""
    normalized_prompt = " ".join(_SYSTEM_PROMPT.split())
    assert "one entry per *section* you touched" in normalized_prompt
    assert "keywords_skipped" in normalized_prompt
    assert '"impact": "high" | "medium" | "low"' in normalized_prompt


def test_prompt_gives_length_guidance_for_every_block_type():
    """Regression guard: unbounded growth in any block risks pushing the
    whole resume past the 2-page budget, where resume_page_fitter.py trims an
    entire entry rather than one bullet — so a future prompt edit must not
    silently drop this length guidance."""
    normalized_prompt = " ".join(_SYSTEM_PROMPT.split())
    assert "roughly the same length" in normalized_prompt  # summary
    assert "compactly" in normalized_prompt  # skills
    assert "1-2 lines" in normalized_prompt  # experience bullets


def test_prompt_instructs_bolding_key_technologies_and_metrics():
    """Regression guard: the emphasis filter (**text** -> <strong>) only has
    an effect if the prompt actually tells the model to use it."""
    assert "**double asterisks**" in _SYSTEM_PROMPT


def test_prompt_explains_soft_skill_and_process_keywords_can_be_truthfully_implied():
    """Regression guard: without this, the model tends to only surface
    literal technology-name keywords and skip general/process ones (e.g.
    "problem solving", "full lifecycle development", "data science") even
    when genuinely implied by real project work."""
    normalized_prompt = " ".join(_SYSTEM_PROMPT.split())
    assert "full lifecycle development" in normalized_prompt
    assert "data science" in normalized_prompt


def test_prompt_explains_the_patch_contract():
    """Regression guard: a future prompt edit must not silently drop the
    block_id/patches contract the parser expects. Structural safety (entry
    count/order, title/company/dates) is now enforced by code — see
    test_synthetic_profile_layout.py — rather than by prompt wording, since
    the LLM has no channel to touch those fields at all."""
    assert "block_id" in _SYSTEM_PROMPT
    assert "editable_blocks" in _SYSTEM_PROMPT


# ---------------------------------------------------------------------------
# Invalid LLM output
# ---------------------------------------------------------------------------

@pytest.mark.anyio
async def test_raises_llm_output_error_on_invalid_json(mock_genai):
    mock_genai.aio.models.generate_content.return_value = MagicMock(text="not json")

    with pytest.raises(LLMOutputError):
        await ResumeGenerationAgent().generate(_PROFILE, _JOB)


@pytest.mark.anyio
async def test_raises_llm_output_error_on_wrong_shaped_json(mock_genai):
    mock_genai.aio.models.generate_content.return_value = _make_response({"skills": "not-a-list"})

    with pytest.raises(LLMOutputError):
        await ResumeGenerationAgent().generate(_PROFILE, _JOB)
