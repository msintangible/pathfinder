import io

import docx

from services.docx_resume_renderer import render_docx


def _make_docx(paragraphs: list[tuple[str, dict]]) -> bytes:
    """paragraphs: list of (text, run_kwargs) e.g. [("hello", {"bold": True})]."""
    document = docx.Document()
    for text, run_kwargs in paragraphs:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        for attr, value in run_kwargs.items():
            setattr(run, attr, value)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _read_paragraphs(docx_bytes: bytes) -> list[str]:
    document = docx.Document(io.BytesIO(docx_bytes))
    return [p.text for p in document.paragraphs]


def test_replaces_a_matched_bullet_in_place():
    source = _make_docx([("Built scalable APIs with Flask and MySQL.", {})])
    original_profile = {"work_experience": [{"bullets": ["Built scalable APIs with Flask and MySQL."]}]}
    optimized_resume = {"experience": [{"bullets": ["Built scalable APIs with Python and AWS."]}]}

    result = render_docx(source, original_profile, optimized_resume)

    assert _read_paragraphs(result) == ["Built scalable APIs with Python and AWS."]


def test_preserves_run_formatting_on_replacement():
    source = _make_docx([("Built scalable APIs with Flask and MySQL.", {"bold": True})])
    original_profile = {"work_experience": [{"bullets": ["Built scalable APIs with Flask and MySQL."]}]}
    optimized_resume = {"experience": [{"bullets": ["Built scalable APIs with Python and AWS."]}]}

    result = render_docx(source, original_profile, optimized_resume)

    document = docx.Document(io.BytesIO(result))
    assert document.paragraphs[0].runs[0].bold is True


def test_leaves_unrelated_paragraphs_untouched():
    source = _make_docx([
        ("Built scalable APIs with Flask and MySQL.", {}),
        ("Education: BSc Computer Science", {}),
    ])
    original_profile = {"work_experience": [{"bullets": ["Built scalable APIs with Flask and MySQL."]}]}
    optimized_resume = {"experience": [{"bullets": ["Built scalable APIs with Python and AWS."]}]}

    result = render_docx(source, original_profile, optimized_resume)

    assert _read_paragraphs(result) == [
        "Built scalable APIs with Python and AWS.",
        "Education: BSc Computer Science",
    ]


def test_no_confident_match_leaves_paragraph_unchanged():
    """The candidate hand-edited this bullet after import, drifting it from the source document."""
    source = _make_docx([("Completely different sentence about something else entirely.", {})])
    original_profile = {"work_experience": [{"bullets": ["Built scalable APIs with Flask."]}]}
    optimized_resume = {"experience": [{"bullets": ["Built scalable APIs with Python."]}]}

    result = render_docx(source, original_profile, optimized_resume)

    assert _read_paragraphs(result) == ["Completely different sentence about something else entirely."]


def test_replaces_summary_and_headline():
    source = _make_docx([
        ("Experienced backend engineer.", {}),
        ("Senior Software Engineer", {}),
    ])
    original_profile = {"summary": "Experienced backend engineer.", "headline": "Senior Software Engineer"}
    optimized_resume = {"summary": "Backend engineer with Python and AWS experience.", "headline": "Senior Backend Engineer"}

    result = render_docx(source, original_profile, optimized_resume)

    assert _read_paragraphs(result) == [
        "Backend engineer with Python and AWS experience.",
        "Senior Backend Engineer",
    ]


def test_replaces_project_description():
    source = _make_docx([("A side project for tracking expenses.", {})])
    original_profile = {"projects": [{"description": "A side project for tracking expenses."}]}
    optimized_resume = {"projects": [{"description": "A full-stack expense tracker built with React and Node.js."}]}

    result = render_docx(source, original_profile, optimized_resume)

    assert _read_paragraphs(result) == ["A full-stack expense tracker built with React and Node.js."]


def test_multi_run_paragraph_collapses_to_a_single_styled_run():
    document = docx.Document()
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run("Built scalable APIs ")
    run1.bold = True
    paragraph.add_run("with Flask and MySQL.")
    buffer = io.BytesIO()
    document.save(buffer)
    source = buffer.getvalue()

    original_profile = {"work_experience": [{"bullets": ["Built scalable APIs with Flask and MySQL."]}]}
    optimized_resume = {"experience": [{"bullets": ["Built scalable APIs with Python and AWS."]}]}

    result = render_docx(source, original_profile, optimized_resume)

    document = docx.Document(io.BytesIO(result))
    assert document.paragraphs[0].text == "Built scalable APIs with Python and AWS."
    assert document.paragraphs[0].runs[0].bold is True
