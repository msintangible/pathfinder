import io

import docx

from schemas.resume_layout import LayoutSection, ResumeLayoutDocument, RunSpan, SectionRole, TextBlock
from services.resume_section_order import DEFAULT_SECTION_ORDER, infer_section_order


def _make_docx(paragraphs: list[tuple[str, str | None]]) -> dict:
    """paragraphs: list of (text, style_name). Returns a real layout_document dict."""
    from services.docx_layout_extractor import extract_docx_layout

    document = docx.Document()
    for text, style_name in paragraphs:
        document.add_paragraph(text, style=style_name)
    buffer = io.BytesIO()
    document.save(buffer)
    return extract_docx_layout(buffer.getvalue()).model_dump()


def _role_section(role: SectionRole) -> LayoutSection:
    return LayoutSection(section_id=f"s-{role.value}", role=role, blocks=[
        TextBlock(block_id=f"{role.value}[0]", kind="paragraph", text="content", runs=[RunSpan(text="content")]),
    ])


def test_defaults_to_standard_order_when_no_layout_document():
    assert infer_section_order(None) == DEFAULT_SECTION_ORDER
    assert infer_section_order({}) == DEFAULT_SECTION_ORDER


def test_invalid_layout_document_falls_back_to_default():
    assert infer_section_order({"not": "a valid layout document"}) == DEFAULT_SECTION_ORDER


def test_infers_order_from_docx_headings_in_document_order():
    layout_document = _make_docx([
        ("Skills", "Heading 1"),
        ("Python, Django", None),
        ("Summary", "Heading 1"),
        ("Backend engineer.", None),
        ("Experience", "Heading 1"),
        ("Built APIs.", None),
    ])

    order = infer_section_order(layout_document)

    assert order == ["skills", "summary", "experience", "projects"]


def test_unrelated_headings_are_ignored():
    layout_document = _make_docx([
        ("Education", "Heading 1"),
        ("BSc Computer Science", None),
        ("Skills", "Heading 1"),
        ("Python", None),
    ])

    order = infer_section_order(layout_document)

    assert order[0] == "skills"
    assert "education" not in order


def test_missing_categories_are_appended_in_default_order():
    layout_document = _make_docx([
        ("Skills", "Heading 1"),
        ("Python, Django", None),
    ])

    order = infer_section_order(layout_document)

    assert order == ["skills", "summary", "experience", "projects"]


def test_infers_order_from_pdf_role_labels():
    layout = ResumeLayoutDocument(source_format="pdf", sections=[
        _role_section(SectionRole.SKILLS),
        _role_section(SectionRole.WORK_EXPERIENCE_ENTRY),
        _role_section(SectionRole.SUMMARY),
    ])

    order = infer_section_order(layout.model_dump())

    assert order == ["skills", "experience", "summary", "projects"]


def test_duplicate_role_sections_only_counted_once():
    layout = ResumeLayoutDocument(source_format="pdf", sections=[
        _role_section(SectionRole.WORK_EXPERIENCE_ENTRY),
        _role_section(SectionRole.WORK_EXPERIENCE_ENTRY),
        _role_section(SectionRole.SKILLS),
    ])

    order = infer_section_order(layout.model_dump())

    assert order == ["experience", "skills", "summary", "projects"]
