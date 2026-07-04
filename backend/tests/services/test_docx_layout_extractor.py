import io

import docx

from services.docx_layout_extractor import extract_docx_layout


def _make_docx(paragraphs: list[tuple[str, str | None, dict]]) -> bytes:
    """paragraphs: list of (text, style_name, run_kwargs), e.g.
    [("Experience", "Heading 1", {}), ("Built APIs.", None, {"bold": True})]."""
    document = docx.Document()
    for text, style_name, run_kwargs in paragraphs:
        paragraph = document.add_paragraph(style=style_name)
        run = paragraph.add_run(text)
        for attr, value in run_kwargs.items():
            setattr(run, attr, value)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _all_blocks(layout):
    return [block for section in layout.sections for block in section.blocks]


def test_source_format_is_docx():
    source = _make_docx([("Hello", None, {})])
    layout = extract_docx_layout(source)
    assert layout.source_format == "docx"


def test_walks_body_paragraphs_and_assigns_stable_block_ids():
    source = _make_docx([("First line.", None, {}), ("Second line.", None, {})])
    layout = extract_docx_layout(source)
    blocks = _all_blocks(layout)
    assert [b.block_id for b in blocks] == ["paragraph[0]", "paragraph[1]"]
    assert [b.text for b in blocks] == ["First line.", "Second line."]


def test_skips_empty_paragraphs():
    source = _make_docx([("First line.", None, {}), ("", None, {}), ("Second line.", None, {})])
    layout = extract_docx_layout(source)
    blocks = _all_blocks(layout)
    assert [b.text for b in blocks] == ["First line.", "Second line."]


def test_splits_sections_on_heading_style():
    source = _make_docx([
        ("Experience", "Heading 1", {}),
        ("Built scalable APIs.", None, {}),
        ("Education", "Heading 1", {}),
        ("BSc Computer Science.", None, {}),
    ])
    layout = extract_docx_layout(source)

    assert len(layout.sections) == 2
    assert [b.text for b in layout.sections[0].blocks] == ["Experience", "Built scalable APIs."]
    assert [b.text for b in layout.sections[1].blocks] == ["Education", "BSc Computer Science."]


def test_content_before_first_heading_is_kept_in_its_own_section():
    source = _make_docx([
        ("Jane Doe — Backend Engineer", None, {}),
        ("Experience", "Heading 1", {}),
        ("Built scalable APIs.", None, {}),
    ])
    layout = extract_docx_layout(source)

    assert len(layout.sections) == 2
    assert layout.sections[0].blocks[0].text == "Jane Doe — Backend Engineer"


def test_bullet_style_paragraph_is_classified_as_bullet():
    source = _make_docx([("Built scalable APIs.", "List Bullet", {})])
    layout = extract_docx_layout(source)
    assert _all_blocks(layout)[0].kind == "bullet"


def test_plain_paragraph_is_classified_as_paragraph():
    source = _make_docx([("Just a sentence.", None, {})])
    layout = extract_docx_layout(source)
    assert _all_blocks(layout)[0].kind == "paragraph"


def test_captures_run_formatting():
    source = _make_docx([("Built scalable APIs.", None, {"bold": True})])
    layout = extract_docx_layout(source)
    runs = _all_blocks(layout)[0].runs
    assert len(runs) == 1
    assert runs[0].text == "Built scalable APIs."
    assert runs[0].bold is True


def test_walks_table_cells_as_their_own_section():
    document = docx.Document()
    document.add_paragraph("Intro paragraph.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "PostgreSQL"
    buffer = io.BytesIO()
    document.save(buffer)
    source = buffer.getvalue()

    layout = extract_docx_layout(source)

    table_sections = [s for s in layout.sections if s.section_id.startswith("table_section")]
    assert len(table_sections) == 1
    table_blocks = table_sections[0].blocks
    assert [b.text for b in table_blocks] == ["Python", "PostgreSQL"]
    assert all(b.kind == "table_cell" for b in table_blocks)


def test_table_cell_block_ids_encode_position_and_are_unique():
    document = docx.Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "Django"
    table.cell(1, 0).text = "AWS"
    table.cell(1, 1).text = "Docker"
    buffer = io.BytesIO()
    document.save(buffer)
    source = buffer.getvalue()

    layout = extract_docx_layout(source)
    blocks = _all_blocks(layout)

    assert [b.block_id for b in blocks] == [
        "table[0].row[0].col[0].paragraph[0]",
        "table[0].row[0].col[1].paragraph[0]",
        "table[0].row[1].col[0].paragraph[0]",
        "table[0].row[1].col[1].paragraph[0]",
    ]
    assert len({b.block_id for b in blocks}) == len(blocks)


def test_block_ids_are_unique_across_paragraphs_and_tables():
    document = docx.Document()
    document.add_paragraph("Intro paragraph.")
    document.add_paragraph("Second paragraph.")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Python"
    buffer = io.BytesIO()
    document.save(buffer)
    source = buffer.getvalue()

    layout = extract_docx_layout(source)
    blocks = _all_blocks(layout)

    assert len({b.block_id for b in blocks}) == len(blocks)


def test_docx_anchor_records_paragraph_index_and_style():
    source = _make_docx([("Experience", "Heading 1", {})])
    layout = extract_docx_layout(source)
    anchor = _all_blocks(layout)[0].docx_anchor

    assert anchor.paragraph_index == 0
    assert anchor.style_name == "Heading 1"
    assert anchor.table_index is None
