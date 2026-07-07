import io

import docx
import docx.opc.constants
import docx.oxml
import pytest
from docx.oxml.ns import qn

from services.docx_layout_extractor import extract_docx_layout
from services.docx_renderer_v2 import DocxRenderError, render_docx
from services.patch_engine import apply_patches
from schemas.resume_layout import ContentPatch, ResumeLayoutDocument


def _add_hyperlink_run(paragraph, text: str, url: str) -> None:
    """Appends a real w:hyperlink run (not a plain Run) to paragraph — mirrors
    how MS Word saves a linked word (e.g. a resume's "GitHub" or portfolio
    URL), which python-docx has no built-in helper for."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = docx.oxml.OxmlElement("w:r")
    run.append(docx.oxml.OxmlElement("w:rPr"))
    t = docx.oxml.OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _save(document: docx.Document) -> bytes:
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _reopen(rendered_bytes: bytes) -> docx.Document:
    return docx.Document(io.BytesIO(rendered_bytes))


def _all_blocks(layout: ResumeLayoutDocument):
    return [block for section in layout.sections for block in section.blocks]


def _block_by_id(layout: ResumeLayoutDocument, block_id: str):
    for block in _all_blocks(layout):
        if block.block_id == block_id:
            return block
    raise AssertionError(f"block_id not found: {block_id}")


def test_replaces_single_run_paragraph_text_and_keeps_formatting():
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Built APIs.").bold = True
    source_bytes = _save(document)

    layout = extract_docx_layout(source_bytes)
    _block_by_id(layout, "paragraph[0]").runs[0].text = "Built scalable APIs."

    rendered = render_docx(source_bytes, layout)
    result_paragraph = _reopen(rendered).paragraphs[0]

    assert result_paragraph.text == "Built scalable APIs."
    assert result_paragraph.runs[0].bold is True


def test_multi_run_paragraph_writes_each_run_independently():
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Built ").bold = True
    paragraph.add_run("scalable APIs.").bold = False
    source_bytes = _save(document)

    layout = extract_docx_layout(source_bytes)
    block = _block_by_id(layout, "paragraph[0]")
    block.runs[0].text = "Led "
    block.runs[1].text = "the migration."

    rendered = render_docx(source_bytes, layout)
    result_runs = _reopen(rendered).paragraphs[0].runs

    assert result_runs[0].text == "Led "
    assert result_runs[0].bold is True
    assert result_runs[1].text == "the migration."
    assert result_runs[1].bold is False


def test_table_cell_text_is_written_in_place():
    document = docx.Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "PostgreSQL"
    source_bytes = _save(document)

    layout = extract_docx_layout(source_bytes)
    block = _block_by_id(layout, "table[0].row[0].col[0].paragraph[0]")
    block.runs[0].text = "Go"

    rendered = render_docx(source_bytes, layout)
    result_table = _reopen(rendered).tables[0]

    assert result_table.cell(0, 0).text == "Go"
    assert result_table.cell(0, 1).text == "PostgreSQL"  # untouched cell is unaffected


def test_paragraph_style_is_left_untouched():
    document = docx.Document()
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.add_run("Experience")
    source_bytes = _save(document)

    layout = extract_docx_layout(source_bytes)
    _block_by_id(layout, "paragraph[0]").runs[0].text = "Professional Experience"

    rendered = render_docx(source_bytes, layout)
    result_paragraph = _reopen(rendered).paragraphs[0]

    assert result_paragraph.text == "Professional Experience"
    assert result_paragraph.style.name == "Heading 1"


def test_blocks_left_unmodified_do_not_change_the_document():
    document = docx.Document()
    document.add_paragraph("First line.")
    document.add_paragraph("Second line.")
    source_bytes = _save(document)

    layout = extract_docx_layout(source_bytes)
    _block_by_id(layout, "paragraph[0]").runs[0].text = "Updated first line."
    # paragraph[1] is left as extracted, unpatched.

    rendered = render_docx(source_bytes, layout)
    result_paragraphs = _reopen(rendered).paragraphs

    assert result_paragraphs[0].text == "Updated first line."
    assert result_paragraphs[1].text == "Second line."


def test_interleaved_empty_run_does_not_break_run_alignment():
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Built ").bold = True
    paragraph.add_run("").italic = True  # empty run the extractor skips over
    paragraph.add_run("scalable APIs.").bold = False
    source_bytes = _save(document)

    layout = extract_docx_layout(source_bytes)
    block = _block_by_id(layout, "paragraph[0]")
    assert len(block.runs) == 2  # confirms the empty run was excluded, as docx_layout_extractor promises
    block.runs[0].text = "Led "
    block.runs[1].text = "the initiative."

    rendered = render_docx(source_bytes, layout)
    result_runs = _reopen(rendered).paragraphs[0].runs

    assert result_runs[0].text == "Led "
    assert result_runs[1].text == ""  # the empty run stays empty, and alignment isn't shifted by it
    assert result_runs[2].text == "the initiative."


def test_run_count_mismatch_falls_back_to_writing_first_run():
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Built ").bold = True
    paragraph.add_run("scalable APIs.").bold = False
    source_bytes = _save(document)

    layout = extract_docx_layout(source_bytes)
    block = _block_by_id(layout, "paragraph[0]")
    # Simulate an inconsistent patched layout (one RunSpan for two real runs)
    # rather than the two the document actually has.
    block.runs = [block.runs[0].model_copy(update={"text": "Completely rewritten bullet."})]

    rendered = render_docx(source_bytes, layout)
    result_runs = _reopen(rendered).paragraphs[0].runs

    assert result_runs[0].text == "Completely rewritten bullet."
    assert result_runs[1].text == ""


def test_hyperlink_text_is_not_duplicated_when_the_surrounding_bullet_is_rewritten():
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Built APIs for ")
    _add_hyperlink_run(paragraph, "GitHub", "https://github.com/example")
    paragraph.add_run(" using Python.")
    source_bytes = _save(document)

    layout = extract_docx_layout(source_bytes)
    patch_result = apply_patches(
        layout, [ContentPatch(block_id="paragraph[0]", new_text="Engineered RESTful APIs using Python and AWS.")]
    )

    rendered = render_docx(source_bytes, patch_result.document)
    result_paragraph = _reopen(rendered).paragraphs[0]

    # The hyperlink's own text ("GitHub") must appear exactly once — before
    # this fix, paragraph.runs was blind to the hyperlink-wrapped run, so the
    # rewrite landed in the surrounding runs while "GitHub" stayed behind
    # untouched, duplicating content instead of replacing it. The pinned
    # hyperlink text still sits at its original position (making the wording
    # around it grammatically seamless is a prompt-level concern, not this
    # mechanical fix's job) — so the new wording surrounds it unchanged.
    assert result_paragraph.text.count("GitHub") == 1
    non_hyperlink_words = result_paragraph.text.replace("GitHub", " ").split()
    assert non_hyperlink_words == "Engineered RESTful APIs using Python and AWS.".split()


def test_hyperlink_relationship_survives_the_rewrite():
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Built APIs for ")
    _add_hyperlink_run(paragraph, "GitHub", "https://github.com/example")
    paragraph.add_run(" using Python.")
    source_bytes = _save(document)

    layout = extract_docx_layout(source_bytes)
    patch_result = apply_patches(
        layout, [ContentPatch(block_id="paragraph[0]", new_text="Engineered RESTful APIs using Python and AWS.")]
    )

    rendered = render_docx(source_bytes, patch_result.document)
    result_paragraph = _reopen(rendered).paragraphs[0]

    hyperlinks = [item for item in result_paragraph.iter_inner_content() if hasattr(item, "address")]
    assert len(hyperlinks) == 1
    assert hyperlinks[0].address == "https://github.com/example"
    assert hyperlinks[0].text == "GitHub"


def test_rejects_non_docx_layout():
    layout = ResumeLayoutDocument(source_format="pdf", sections=[])
    document = docx.Document()

    with pytest.raises(DocxRenderError):
        render_docx(_save(document), layout)
