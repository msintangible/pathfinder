import io

import docx

from schemas.resume_layout import LayoutSection, PdfAnchor, ResumeLayoutDocument, RunSpan, SectionRole, TextBlock
from services.docx_layout_extractor import extract_docx_layout
from services.profile_layout_correlator import correlate_profile_to_layout


def _make_docx(paragraphs: list[tuple[str, str | None]]) -> bytes:
    """paragraphs: list of (text, style_name)."""
    document = docx.Document()
    for text, style_name in paragraphs:
        document.add_paragraph(text, style=style_name)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_PROFILE = {
    "headline": "Senior Backend Engineer",
    "summary": "Experienced backend engineer specialising in Python and AWS.",
    "work_experience": [
        {"bullets": ["Built scalable APIs with Python and Django"]},
    ],
    "projects": [
        {"description": "An AI-powered job application assistant"},
    ],
}


def _docx_layout(paragraphs: list[tuple[str, str | None]]) -> ResumeLayoutDocument:
    return extract_docx_layout(_make_docx(paragraphs))


# ---------------------------------------------------------------------------
# Exact / near-exact matching against a docx layout
# ---------------------------------------------------------------------------

def test_matches_headline_summary_and_bullet_to_real_blocks():
    layout = _docx_layout([
        ("Senior Backend Engineer", None),
        ("Experienced backend engineer specialising in Python and AWS.", None),
        ("Built scalable APIs with Python and Django", None),
        ("An AI-powered job application assistant", None),
    ])

    result = correlate_profile_to_layout(_PROFILE, layout)

    assert result.block_id_map["headline"] == "paragraph[0]"
    assert result.block_id_map["summary"] == "paragraph[1]"
    assert result.block_id_map["work_experience[0].bullets[0]"] == "paragraph[2]"
    assert result.block_id_map["projects[0].description"] == "paragraph[3]"
    assert result.match_rate == 1.0


def test_matches_despite_minor_whitespace_and_casing_differences():
    # The real document has slightly different whitespace/casing than the
    # profile's extracted text — this is expected minor extraction drift,
    # not a wording change, and should still clear the similarity threshold.
    layout = _docx_layout([("  built SCALABLE apis with python and django  ", None)])
    profile = {"work_experience": [{"bullets": ["Built scalable APIs with Python and Django"]}]}

    result = correlate_profile_to_layout(profile, layout)

    assert result.block_id_map["work_experience[0].bullets[0]"] == "paragraph[0]"
    assert result.match_rate == 1.0


def test_leaves_dissimilar_text_uncorrelated():
    layout = _docx_layout([("Something completely unrelated to any profile field", None)])
    profile = {"work_experience": [{"bullets": ["Built scalable APIs with Python and Django"]}]}

    result = correlate_profile_to_layout(profile, layout)

    assert "work_experience[0].bullets[0]" not in result.block_id_map
    assert result.match_rate == 0.0


def test_match_rate_reflects_partial_correlation():
    layout = _docx_layout([("Built scalable APIs with Python and Django", None)])  # only one of two bullets present
    profile = {
        "work_experience": [
            {"bullets": ["Built scalable APIs with Python and Django", "A bullet with no document counterpart"]}
        ]
    }

    result = correlate_profile_to_layout(profile, layout)

    assert result.match_rate == 0.5


def test_each_real_block_is_used_at_most_once():
    # Two profile bullets share identical text — each must resolve to a
    # *different* real paragraph, not both claiming the same one.
    layout = _docx_layout([
        ("Wrote unit tests", None),
        ("Wrote unit tests", None),
    ])
    profile = {"work_experience": [{"bullets": ["Wrote unit tests", "Wrote unit tests"]}]}

    result = correlate_profile_to_layout(profile, layout)

    matched_ids = {result.block_id_map["work_experience[0].bullets[0]"], result.block_id_map["work_experience[0].bullets[1]"]}
    assert matched_ids == {"paragraph[0]", "paragraph[1]"}


def test_empty_profile_has_zero_match_rate_and_empty_map():
    layout = _docx_layout([("Some paragraph.", None)])

    result = correlate_profile_to_layout({}, layout)

    assert result.match_rate == 0.0
    assert result.block_id_map == {}


def test_global_matching_resolves_declaration_order_contention():
    # bullets[0] is declared first and is a decent (not perfect) match for
    # paragraph[0] — its only qualifying candidate. bullets[1] is an *exact*
    # match for that same paragraph[0] and has no other real-document
    # counterpart at all (its similarity to paragraph[1] is well below
    # threshold). Field-declaration-order greedy would let bullets[0] claim
    # paragraph[0] first since it's bullets[0]'s best available option,
    # leaving bullets[1] unmatched even though it's a far stronger candidate
    # for that same block. Global best-match-first assignment must instead
    # give paragraph[0] to bullets[1] (the single strongest pair overall) and
    # fall bullets[0] back to its second-best option, paragraph[1].
    paragraph_0 = "alpha bravo charlie delta echo foxtrot golf hotel india juliet"
    paragraph_1 = "kilo lima mike november oscar papa quebec romeo sierra tango"
    bullet_0 = "alpha bravo charlie delta echo foxtrot golf kilo lima mike november oscar papa quebec"
    bullet_1 = paragraph_0

    layout = _docx_layout([(paragraph_0, None), (paragraph_1, None)])
    profile = {"work_experience": [{"bullets": [bullet_0, bullet_1]}]}

    result = correlate_profile_to_layout(profile, layout)

    assert result.block_id_map["work_experience[0].bullets[1]"] == "paragraph[0]"
    assert result.block_id_map["work_experience[0].bullets[0]"] == "paragraph[1]"
    assert result.match_rate == 1.0


# ---------------------------------------------------------------------------
# Skills section detection — docx heading heuristic
# ---------------------------------------------------------------------------

def test_finds_skills_block_via_docx_heading_keyword():
    layout = _docx_layout([
        ("Skills", "Heading 1"),
        ("Python, Django, PostgreSQL", None),
    ])

    result = correlate_profile_to_layout({}, layout)

    assert result.block_id_map["skills"] == "paragraph[1]"
    assert result.skills_overflow_block_ids == []


def test_skills_field_is_excluded_from_match_rate_denominator():
    layout = _docx_layout([
        ("Skills", "Heading 1"),
        ("Python, Django, PostgreSQL", None),
    ])

    result = correlate_profile_to_layout({}, layout)

    assert result.match_rate == 0.0  # no correlatable (non-skills) fields at all, not skewed by the skills match


def test_multi_block_skills_section_uses_first_block_and_flags_overflow():
    layout = _docx_layout([
        ("Skills", "Heading 1"),
        ("Python", None),
        ("PostgreSQL", None),
    ])

    result = correlate_profile_to_layout({}, layout)

    assert result.block_id_map["skills"] == "paragraph[1]"
    assert result.skills_overflow_block_ids == ["paragraph[2]"]


def test_three_block_skills_section_flags_all_but_first_as_overflow():
    layout = _docx_layout([
        ("Skills", "Heading 1"),
        ("Languages: Python, Java", None),
        ("Frameworks: Django, React", None),
        ("Tools: Docker, Git", None),
    ])

    result = correlate_profile_to_layout({}, layout)

    assert result.block_id_map["skills"] == "paragraph[1]"
    assert result.skills_overflow_block_ids == ["paragraph[2]", "paragraph[3]"]


def test_non_skills_headings_are_not_mistaken_for_a_skills_section():
    layout = _docx_layout([
        ("Education", "Heading 1"),
        ("BSc Computer Science", None),
    ])

    result = correlate_profile_to_layout({}, layout)

    assert "skills" not in result.block_id_map


# ---------------------------------------------------------------------------
# Skills section detection — pdf role label (from gemini_vision_layout_agent.py)
# ---------------------------------------------------------------------------

def _pdf_block(block_id: str, text: str) -> TextBlock:
    return TextBlock(
        block_id=block_id, kind="paragraph", text=text, runs=[RunSpan(text=text)],
        pdf_anchor=PdfAnchor(page_number=0, x0=0, y0=0, x1=100, y1=10),
    )


def _pdf_heading_block(block_id: str, text: str) -> TextBlock:
    """A real PDF section header: one run, entirely bold, short text."""
    return TextBlock(
        block_id=block_id, kind="paragraph", text=text, runs=[RunSpan(text=text, bold=True)],
        pdf_anchor=PdfAnchor(page_number=0, x0=0, y0=0, x1=100, y1=10),
    )


def _pdf_labeled_content_block(block_id: str, label: str, rest: str) -> TextBlock:
    """A content line with a bold label prefix and plain body — e.g.
    "**Developer Tools**: AWS, Postman, ..." — must NOT be mistaken for a
    heading even though it starts with a bold run."""
    return TextBlock(
        block_id=block_id, kind="paragraph", text=f"{label}{rest}",
        runs=[RunSpan(text=label, bold=True), RunSpan(text=rest, bold=False)],
        pdf_anchor=PdfAnchor(page_number=0, x0=0, y0=0, x1=100, y1=10),
    )


def test_finds_skills_block_via_pdf_role_label():
    layout = ResumeLayoutDocument(source_format="pdf", sections=[
        LayoutSection(section_id="s0", role=SectionRole.SKILLS, blocks=[
            _pdf_block("page[0].block[0].line[0]", "Python, Django, PostgreSQL"),
        ]),
    ])

    result = correlate_profile_to_layout({}, layout)

    assert result.block_id_map["skills"] == "page[0].block[0].line[0]"


def test_pdf_skills_heading_is_excluded_from_content_blocks():
    # Mirrors a real multi-line PDF skills section: a bold "Technical
    # Skills" heading followed by category lines that each start with a
    # bold label. The heading must not be picked as the primary skills
    # block, and the labeled content lines must not be mistaken for headings
    # just because they start with a bold run.
    layout = ResumeLayoutDocument(source_format="pdf", sections=[
        LayoutSection(section_id="s0", role=SectionRole.SKILLS, blocks=[
            _pdf_heading_block("page[0].block[0]", "Technical Skills"),
            _pdf_labeled_content_block("page[0].block[1]", "Languages: ", "Python, Java, SQL"),
            _pdf_labeled_content_block("page[0].block[2]", "Developer Tools: ", "AWS, Postman, Git"),
        ]),
    ])

    result = correlate_profile_to_layout({}, layout)

    assert result.block_id_map["skills"] == "page[0].block[1]"
    assert result.skills_overflow_block_ids == ["page[0].block[2]"]


def test_pdf_skills_heading_keyword_fallback_now_works():
    # No role assigned (defaults to OTHER, as if vision labeling failed or
    # mislabeled this page) — detection must still succeed via the
    # heading-keyword fallback, now that is_heading_block recognizes PDF
    # headings structurally.
    layout = ResumeLayoutDocument(source_format="pdf", sections=[
        LayoutSection(section_id="s0", blocks=[
            _pdf_heading_block("page[0].block[0]", "Technical Skills"),
            _pdf_block("page[0].block[1]", "Python, Django, PostgreSQL"),
        ]),
    ])

    result = correlate_profile_to_layout({}, layout)

    assert result.block_id_map["skills"] == "page[0].block[1]"
