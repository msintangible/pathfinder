"""
Regression coverage for POST /v1/profile/import's empty-content guard.

Found via live verification: calling CandidateProfileAgent with a blank
prompt (no CV, no LinkedIn text, a GitHub/portfolio fetch that both came back
empty) made Gemini fabricate an entire fictional profile instead of returning
nulls, despite the agent's "never invent" system instruction. The endpoint
must refuse before ever reaching the agent in that case.
"""
import io
import uuid
from unittest.mock import AsyncMock, patch

import docx
import fitz
import pytest
from fastapi import FastAPI
from fastapi.testclient import TestClient

from api.v1.profile import router
from core.security import get_current_user
from database.session import get_db
from models.profile import UserProfile
from models.user import User
from schemas.resume_layout import LayoutSection, ResumeLayoutDocument, SectionRole, TextBlock

_TEST_USER_ID = uuid.uuid4()


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _make_pdf_bytes(lines: list[str]) -> bytes:
    document = fitz.open()
    page = document.new_page()
    for index, text in enumerate(lines):
        page.insert_text((72, 100 + index * 20), text, fontsize=11)
    buffer = document.tobytes()
    document.close()
    return buffer


@pytest.fixture
def client():
    app = FastAPI()
    app.include_router(router, prefix="/v1")
    app.dependency_overrides[get_db] = lambda: None
    app.dependency_overrides[get_current_user] = lambda: User(id=_TEST_USER_ID)
    return TestClient(app)


def test_rejects_when_no_content_could_be_extracted(client):
    """A github_url that resolves to nothing, with no other source, must 422 — not fabricate a profile."""
    with patch("api.v1.profile.fetch_github_profile", new=AsyncMock(return_value=(None, []))), \
         patch("api.v1.profile.fetch_portfolio_text", new=AsyncMock(return_value=None)), \
         patch("api.v1.profile.CandidateProfileAgent") as mock_agent_cls:
        mock_agent_cls.return_value.analyze = AsyncMock()

        resp = client.post(
            "/v1/profile/import",
            data={"github_url": "https://github.com/this-user-should-not-exist-xyz123"},
        )

        assert resp.status_code == 422
        mock_agent_cls.return_value.analyze.assert_not_called()


def test_proceeds_when_linkedin_text_present(client):
    """Real scraped content (even with no other source) must reach the agent."""
    with patch("api.v1.profile.fetch_github_profile", new=AsyncMock(return_value=(None, []))), \
         patch("api.v1.profile.fetch_portfolio_text", new=AsyncMock(return_value=None)), \
         patch("api.v1.profile.CandidateProfileAgent") as mock_agent_cls, \
         patch("api.v1.profile.ProfileRepository") as mock_repo_cls:
        mock_agent_cls.return_value.analyze = AsyncMock(return_value={"name": "Jane Doe"})
        test_id = uuid.uuid4()
        mock_repo_cls.return_value.create_from_analysis = AsyncMock(
            return_value=UserProfile(id=test_id, name="Jane Doe")
        )

        resp = client.post(
            "/v1/profile/import",
            data={"linkedin_url": "https://linkedin.com/in/jane", "linkedin_text": "Jane Doe, Engineer"},
        )

        assert resp.status_code == 200
        mock_agent_cls.return_value.analyze.assert_called_once()
        mock_repo_cls.return_value.create_from_analysis.assert_called_once()
        assert mock_repo_cls.return_value.create_from_analysis.call_args.kwargs["user_id"] == _TEST_USER_ID

        body = resp.json()
        assert body["id"] == str(test_id)
        assert body["sources"]["linkedin_text"] == "Jane Doe, Engineer"
        assert body["sources"]["resume_text"] is None
        assert body["sources"]["github_repositories"] == []


def test_docx_upload_stores_the_source_document(client, tmp_path):
    """A .docx CV must be stored so resume generation can later edit it in place, not just have its text extracted."""
    docx_bytes = _make_docx_bytes(["Jane Doe", "Senior Backend Engineer"])

    with patch("api.v1.profile.fetch_github_profile", new=AsyncMock(return_value=(None, []))), \
         patch("api.v1.profile.fetch_portfolio_text", new=AsyncMock(return_value=None)), \
         patch("api.v1.profile.CandidateProfileAgent") as mock_agent_cls, \
         patch("api.v1.profile.ProfileRepository") as mock_repo_cls, \
         patch("api.v1.profile.LocalResumeStorage") as mock_storage_cls:
        mock_agent_cls.return_value.analyze = AsyncMock(return_value={"name": "Jane Doe"})
        mock_repo_cls.return_value.create_from_analysis = AsyncMock(
            return_value=UserProfile(id=uuid.uuid4(), name="Jane Doe")
        )
        stored_path = str(tmp_path / "source-abc123.docx")
        mock_storage_cls.return_value.save.return_value = stored_path

        resp = client.post(
            "/v1/profile/import",
            files={"file": ("resume.docx", docx_bytes,
                             "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )

        assert resp.status_code == 200
        mock_storage_cls.return_value.save.assert_called_once()
        saved_bytes = mock_storage_cls.return_value.save.call_args.args[0]
        assert saved_bytes == docx_bytes

        call_kwargs = mock_repo_cls.return_value.create_from_analysis.call_args.kwargs
        assert call_kwargs["source_document_path"] == stored_path
        assert call_kwargs["source_document_format"] == "docx"


def test_docx_upload_also_builds_a_layout_document(client, tmp_path):
    """The deterministic docx layout extractor must run at import time so
    resume generation has block-level structure to edit later, not just the
    plain concatenated text used for profile analysis."""
    docx_bytes = _make_docx_bytes(["Jane Doe", "Senior Backend Engineer"])

    with patch("api.v1.profile.fetch_github_profile", new=AsyncMock(return_value=(None, []))), \
         patch("api.v1.profile.fetch_portfolio_text", new=AsyncMock(return_value=None)), \
         patch("api.v1.profile.CandidateProfileAgent") as mock_agent_cls, \
         patch("api.v1.profile.ProfileRepository") as mock_repo_cls, \
         patch("api.v1.profile.LocalResumeStorage") as mock_storage_cls:
        mock_agent_cls.return_value.analyze = AsyncMock(return_value={"name": "Jane Doe"})
        mock_repo_cls.return_value.create_from_analysis = AsyncMock(
            return_value=UserProfile(id=uuid.uuid4(), name="Jane Doe")
        )
        mock_storage_cls.return_value.save.return_value = str(tmp_path / "source-abc123.docx")

        resp = client.post(
            "/v1/profile/import",
            files={"file": ("resume.docx", docx_bytes,
                             "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )

        assert resp.status_code == 200
        layout_document = mock_repo_cls.return_value.create_from_analysis.call_args.kwargs["layout_document"]
        assert layout_document["source_format"] == "docx"
        block_texts = [b["text"] for s in layout_document["sections"] for b in s["blocks"]]
        assert block_texts == ["Jane Doe", "Senior Backend Engineer"]


def test_docx_upload_falls_back_to_no_layout_document_when_extraction_fails(client, tmp_path):
    docx_bytes = _make_docx_bytes(["Jane Doe"])

    with patch("api.v1.profile.fetch_github_profile", new=AsyncMock(return_value=(None, []))), \
         patch("api.v1.profile.fetch_portfolio_text", new=AsyncMock(return_value=None)), \
         patch("api.v1.profile.CandidateProfileAgent") as mock_agent_cls, \
         patch("api.v1.profile.ProfileRepository") as mock_repo_cls, \
         patch("api.v1.profile.LocalResumeStorage") as mock_storage_cls, \
         patch("api.v1.profile.extract_docx_layout") as mock_extract_layout:
        from services.docx_layout_extractor import DocxLayoutExtractionError
        mock_extract_layout.side_effect = DocxLayoutExtractionError("corrupt")
        mock_agent_cls.return_value.analyze = AsyncMock(return_value={"name": "Jane Doe"})
        mock_repo_cls.return_value.create_from_analysis = AsyncMock(
            return_value=UserProfile(id=uuid.uuid4(), name="Jane Doe")
        )
        mock_storage_cls.return_value.save.return_value = str(tmp_path / "source-abc123.docx")

        resp = client.post(
            "/v1/profile/import",
            files={"file": ("resume.docx", docx_bytes,
                             "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )

        assert resp.status_code == 200  # a broken layout extraction must not block the import
        call_kwargs = mock_repo_cls.return_value.create_from_analysis.call_args.kwargs
        assert call_kwargs["layout_document"] is None


def test_pdf_upload_stores_the_source_document_and_layout(client, tmp_path):
    """A .pdf CV must be stored, and get a layout document (extraction + vision labeling), the same as docx."""
    pdf_bytes = _make_pdf_bytes(["Jane Doe", "Senior Backend Engineer"])
    labeled_layout = ResumeLayoutDocument(source_format="pdf", sections=[
        LayoutSection(section_id="page_section[0].labeled[0]", role=SectionRole.HEADER_CONTACT, blocks=[
            TextBlock(block_id="page[0].block[0].line[0]", kind="paragraph", text="Jane Doe"),
        ]),
    ])

    with patch("api.v1.profile.fetch_github_profile", new=AsyncMock(return_value=(None, []))), \
         patch("api.v1.profile.fetch_portfolio_text", new=AsyncMock(return_value=None)), \
         patch("api.v1.profile.CandidateProfileAgent") as mock_agent_cls, \
         patch("api.v1.profile.ProfileRepository") as mock_repo_cls, \
         patch("api.v1.profile.LocalResumeStorage") as mock_storage_cls, \
         patch("api.v1.profile.GeminiVisionLayoutAgent") as mock_vision_cls:
        mock_agent_cls.return_value.analyze = AsyncMock(return_value={"name": "Jane Doe"})
        mock_repo_cls.return_value.create_from_analysis = AsyncMock(
            return_value=UserProfile(id=uuid.uuid4(), name="Jane Doe")
        )
        stored_path = str(tmp_path / "source-abc123.pdf")
        mock_storage_cls.return_value.save.return_value = stored_path
        mock_vision_cls.return_value.label_document = AsyncMock(return_value=labeled_layout)

        resp = client.post(
            "/v1/profile/import",
            files={"file": ("resume.pdf", pdf_bytes, "application/pdf")},
        )

        assert resp.status_code == 200
        mock_storage_cls.return_value.save.assert_called_once()
        assert mock_storage_cls.return_value.save.call_args.args[0] == pdf_bytes

        call_kwargs = mock_repo_cls.return_value.create_from_analysis.call_args.kwargs
        assert call_kwargs["source_document_path"] == stored_path
        assert call_kwargs["source_document_format"] == "pdf"
        assert call_kwargs["layout_document"] == labeled_layout.model_dump()


def test_pdf_layout_document_survives_vision_labeling_failure(client, tmp_path):
    """A Gemini outage/quota error during vision labeling is a best-effort
    enhancement failing, not a reason to fail the whole import — the
    deterministic (unlabeled) extraction must still be persisted."""
    pdf_bytes = _make_pdf_bytes(["Jane Doe"])

    with patch("api.v1.profile.fetch_github_profile", new=AsyncMock(return_value=(None, []))), \
         patch("api.v1.profile.fetch_portfolio_text", new=AsyncMock(return_value=None)), \
         patch("api.v1.profile.CandidateProfileAgent") as mock_agent_cls, \
         patch("api.v1.profile.ProfileRepository") as mock_repo_cls, \
         patch("api.v1.profile.LocalResumeStorage") as mock_storage_cls, \
         patch("api.v1.profile.GeminiVisionLayoutAgent") as mock_vision_cls:
        mock_agent_cls.return_value.analyze = AsyncMock(return_value={"name": "Jane Doe"})
        mock_repo_cls.return_value.create_from_analysis = AsyncMock(
            return_value=UserProfile(id=uuid.uuid4(), name="Jane Doe")
        )
        mock_storage_cls.return_value.save.return_value = str(tmp_path / "source-abc123.pdf")
        mock_vision_cls.return_value.label_document = AsyncMock(side_effect=RuntimeError("Gemini is down"))

        resp = client.post(
            "/v1/profile/import",
            files={"file": ("resume.pdf", pdf_bytes, "application/pdf")},
        )

        assert resp.status_code == 200
        layout_document = mock_repo_cls.return_value.create_from_analysis.call_args.kwargs["layout_document"]
        assert layout_document is not None
        assert layout_document["source_format"] == "pdf"
        assert layout_document["sections"][0]["role"] == "other"  # unlabeled fallback


def test_pdf_upload_falls_back_to_no_layout_document_when_extraction_fails(client, tmp_path):
    with patch("api.v1.profile.fetch_github_profile", new=AsyncMock(return_value=(None, []))), \
         patch("api.v1.profile.fetch_portfolio_text", new=AsyncMock(return_value=None)), \
         patch("api.v1.profile.CandidateProfileAgent") as mock_agent_cls, \
         patch("api.v1.profile.ProfileRepository") as mock_repo_cls, \
         patch("api.v1.profile.LocalResumeStorage") as mock_storage_cls, \
         patch("api.v1.profile.extract_pdf_layout") as mock_extract_layout, \
         patch("api.v1.profile.extract_pdf_text", return_value="Jane Doe"):
        from services.pdf_layout_extractor import PdfLayoutExtractionError
        mock_extract_layout.side_effect = PdfLayoutExtractionError("corrupt")
        mock_agent_cls.return_value.analyze = AsyncMock(return_value={"name": "Jane Doe"})
        mock_repo_cls.return_value.create_from_analysis = AsyncMock(
            return_value=UserProfile(id=uuid.uuid4(), name="Jane Doe")
        )
        mock_storage_cls.return_value.save.return_value = str(tmp_path / "source-abc123.pdf")

        resp = client.post(
            "/v1/profile/import",
            files={"file": ("resume.pdf", b"%PDF-fake", "application/pdf")},
        )

        assert resp.status_code == 200
        call_kwargs = mock_repo_cls.return_value.create_from_analysis.call_args.kwargs
        assert call_kwargs["layout_document"] is None


def test_rejects_unsupported_file_type(client):
    with patch("api.v1.profile.CandidateProfileAgent") as mock_agent_cls:
        mock_agent_cls.return_value.analyze = AsyncMock()

        resp = client.post(
            "/v1/profile/import",
            files={"file": ("resume.txt", b"plain text resume", "text/plain")},
        )

        assert resp.status_code == 400
        mock_agent_cls.return_value.analyze.assert_not_called()


def test_requires_authentication():
    """Without a valid bearer token, the endpoint must 401 before touching any source."""
    app = FastAPI()
    app.include_router(router, prefix="/v1")
    app.dependency_overrides[get_db] = lambda: None
    unauthenticated_client = TestClient(app)

    resp = unauthenticated_client.post(
        "/v1/profile/import",
        data={"linkedin_url": "https://linkedin.com/in/jane", "linkedin_text": "Jane Doe, Engineer"},
    )

    assert resp.status_code == 401


def test_rejects_malformed_portfolio_url(client):
    """A malformed portfolio_url must 422 before any fetch or agent call."""
    with patch("api.v1.profile.fetch_github_profile", new=AsyncMock(return_value=(None, []))), \
         patch("api.v1.profile.fetch_portfolio_text", new=AsyncMock()) as mock_fetch_portfolio, \
         patch("api.v1.profile.CandidateProfileAgent") as mock_agent_cls:
        mock_agent_cls.return_value.analyze = AsyncMock()

        resp = client.post(
            "/v1/profile/import",
            data={"portfolio_url": "not-a-url"},
        )

        assert resp.status_code == 422
        mock_fetch_portfolio.assert_not_called()
        mock_agent_cls.return_value.analyze.assert_not_called()
