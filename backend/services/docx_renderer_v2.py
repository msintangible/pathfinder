"""
DOCX Renderer — writes an already-patched ResumeLayoutDocument back into the
candidate's original .docx bytes.

Purely mechanical: by the time a ResumeLayoutDocument reaches this module,
docx_layout_extractor.py has already built it from these same bytes and
patch_engine.py has already redistributed every wording change across each
block's existing RunSpans. This module's only job is translating that
already-finalized set of run texts into real python-docx paragraph/table-
cell/run writes, in place — it contains no patch-matching or text-
redistribution logic of its own, and never touches paragraph-level
formatting (pPr/style).

Wired into api/v1/resume.py::generate_resume, replacing the deleted
docx_resume_renderer.py (Phase 9 of the structure-preserving resume plan).
"""

import io
import logging

from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from schemas.resume_layout import DocxAnchor, ResumeLayoutDocument, RunSpan
from services.docx_layout_extractor import iter_run_targets

logger = logging.getLogger(__name__)


class DocxRenderError(Exception):
    pass


def render_docx(original_bytes: bytes, layout: ResumeLayoutDocument) -> bytes:
    if layout.source_format != "docx":
        raise DocxRenderError(f"Expected a docx-sourced layout, got source_format={layout.source_format!r}")

    document = Document(io.BytesIO(original_bytes))
    for section in layout.sections:
        for block in section.blocks:
            if block.docx_anchor is None:
                continue
            paragraph = _resolve_paragraph(document, block.docx_anchor)
            _write_runs(paragraph, block.runs)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _resolve_paragraph(document: Document, anchor: DocxAnchor) -> Paragraph:
    if anchor.table_index is not None:
        cell: _Cell = document.tables[anchor.table_index].rows[anchor.row_index].cells[anchor.col_index]
        return cell.paragraphs[anchor.cell_paragraph_index]
    return document.paragraphs[anchor.paragraph_index]


def _write_runs(paragraph: Paragraph, patched_runs: list[RunSpan]) -> None:
    """
    Write patched_runs' text into paragraph's real runs, in order.

    Uses iter_run_targets (shared with docx_layout_extractor.py) instead of
    paragraph.runs — paragraph.runs silently excludes runs wrapped in a
    w:hyperlink element, which would otherwise leave a hyperlink's original
    text untouched in the document while the rest of the block's new text is
    written around it. patch_engine.py pins hyperlink RunSpans to their
    original text, so writing them back here is a harmless no-op — it's what
    keeps the hyperlink attached to unchanged visible text.
    """
    targets = [run for run, _ in iter_run_targets(paragraph)]

    if len(targets) != len(patched_runs):
        logger.warning(
            "docx run count mismatch (document has %d non-empty runs, patch has %d) — "
            "falling back to writing combined text into the first run",
            len(targets), len(patched_runs),
        )
        _write_fallback(targets, patched_runs)
        return

    for target, patched in zip(targets, patched_runs):
        target.text = patched.text


def _write_fallback(targets: list[Run], patched_runs: list[RunSpan]) -> None:
    if not targets:
        return
    targets[0].text = "".join(patched.text for patched in patched_runs)
    for run in targets[1:]:
        run.text = ""
