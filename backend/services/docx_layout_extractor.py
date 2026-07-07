"""
Builds a ResumeLayoutDocument from an uploaded .docx file.

Walks both document.paragraphs and document.tables — the plain-text
extractor (docx_text_extractor.py) and the original renderer
(docx_resume_renderer.py) only ever look at document.paragraphs, so
table-based resume sections (e.g. a skills grid) were never previously
represented at all.

block_id is a stable structural identifier derived from document position
only (e.g. "paragraph[5]", "table[0].row[2].col[1].paragraph[0]") — it does
not yet know which CandidateProfile field a block corresponds to. That
semantic correlation happens in a later phase, once the profile has been
extracted from this same document's text.
"""

import io
import zipfile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from schemas.resume_layout import DocxAnchor, LayoutSection, ResumeLayoutDocument, RunSpan, TextBlock


class DocxLayoutExtractionError(Exception):
    pass


def _is_heading(paragraph: Paragraph) -> bool:
    style_name = (paragraph.style.name or "") if paragraph.style else ""
    return style_name.lower().startswith("heading") or style_name.lower() == "title"


def _is_bullet_style(style_name: str | None) -> bool:
    return bool(style_name) and "list" in style_name.lower()


def iter_run_targets(paragraph: Paragraph) -> list[tuple[Run, str | None]]:
    """
    Flattens a paragraph's inner content into (run, hyperlink_url) pairs, in
    document order, skipping empty-text runs.

    python-docx's paragraph.runs property silently excludes any run wrapped
    in a w:hyperlink element (e.g. a linked "GitHub" or portfolio URL) — used
    directly, that would make hyperlinked text invisible to both extraction
    and rendering, leaving its original text sitting untouched in the
    document alongside newly-written text for the rest of the block instead
    of actually being accounted for. docx_renderer_v2.py reuses this same
    helper (rather than paragraph.runs) so run positions always line up
    between reading and writing.
    """
    targets: list[tuple[Run, str | None]] = []
    for item in paragraph.iter_inner_content():
        if isinstance(item, Hyperlink):
            for run in item.runs:
                if run.text:
                    targets.append((run, item.address))
        elif item.text:
            targets.append((item, None))
    return targets


def _run_spans(paragraph: Paragraph) -> list[RunSpan]:
    spans = []
    for run, hyperlink_url in iter_run_targets(paragraph):
        spans.append(RunSpan(
            text=run.text,
            bold=bool(run.bold),
            italic=bool(run.italic),
            underline=bool(run.underline),
            font_name=run.font.name,
            font_size=run.font.size.pt if run.font.size else None,
            hyperlink_url=hyperlink_url,
        ))
    return spans


def _extract_paragraph_sections(document: Document) -> list[LayoutSection]:
    """Body paragraphs, split into sections at each heading-styled paragraph."""
    sections: list[LayoutSection] = [LayoutSection(section_id="section[0]")]
    for index, paragraph in enumerate(document.paragraphs):
        if _is_heading(paragraph):
            sections.append(LayoutSection(section_id=f"section[{len(sections)}]"))
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else None
        kind = "bullet" if _is_bullet_style(style_name) else "paragraph"
        sections[-1].blocks.append(TextBlock(
            block_id=f"paragraph[{index}]",
            kind=kind,
            text=text,
            runs=_run_spans(paragraph),
            docx_anchor=DocxAnchor(paragraph_index=index, style_name=style_name),
        ))
    # The leading placeholder section only matters if content preceded the
    # first heading — drop it if the document opens with a heading.
    if len(sections) > 1 and not sections[0].blocks:
        sections.pop(0)
    return sections


def _extract_table_sections(document: Document) -> list[LayoutSection]:
    """Each table becomes its own section — resume tables (skills grids,
    timelines) are typically self-contained components, not sub-parts of the
    surrounding paragraph flow."""
    sections: list[LayoutSection] = []
    for table_index, table in enumerate(document.tables):
        section = LayoutSection(section_id=f"table_section[{table_index}]")
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                for para_index, paragraph in enumerate(cell.paragraphs):
                    text = paragraph.text.strip()
                    if not text:
                        continue
                    style_name = paragraph.style.name if paragraph.style else None
                    section.blocks.append(TextBlock(
                        block_id=f"table[{table_index}].row[{row_index}].col[{col_index}].paragraph[{para_index}]",
                        kind="table_cell",
                        text=text,
                        runs=_run_spans(paragraph),
                        docx_anchor=DocxAnchor(
                            table_index=table_index, row_index=row_index,
                            col_index=col_index, cell_paragraph_index=para_index,
                            style_name=style_name,
                        ),
                    ))
        if section.blocks:
            sections.append(section)
    return sections


def extract_docx_layout(docx_bytes: bytes) -> ResumeLayoutDocument:
    try:
        document = Document(io.BytesIO(docx_bytes))
    except (PackageNotFoundError, zipfile.BadZipFile, KeyError, ValueError) as exc:
        raise DocxLayoutExtractionError(f"Could not read DOCX: {exc}") from exc

    sections = _extract_paragraph_sections(document) + _extract_table_sections(document)
    return ResumeLayoutDocument(source_format="docx", sections=sections)
