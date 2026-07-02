import io
from difflib import SequenceMatcher

from docx import Document
from docx.text.paragraph import Paragraph

# Below this similarity, a paragraph is treated as "not the same sentence
# anymore" (e.g. the candidate hand-edited that bullet after import) and is
# left untouched rather than risking a wrong replacement.
_MATCH_THRESHOLD = 0.6


def _collect_replacements(original_profile: dict, optimized_resume: dict) -> list[tuple[str, str]]:
    """
    Pair up (original_text, tailored_text) positionally between the
    pre-tailoring profile and the LLM's output. Relies on
    ResumeGenerationAgent's structure-preserving prompt guaranteeing the same
    number of experience/project entries, in the same order, as given —
    without that, there'd be no reliable way to know which "before" bullet
    corresponds to which "after" bullet.

    Only summary/headline/bullets/descriptions are in scope — skills sections
    vary too structurally to reliably locate as one field, and
    changes_summary isn't part of the original document at all.
    """
    pairs: list[tuple[str, str]] = []

    if original_profile.get("summary") and optimized_resume.get("summary"):
        pairs.append((original_profile["summary"], optimized_resume["summary"]))
    if original_profile.get("headline") and optimized_resume.get("headline"):
        pairs.append((original_profile["headline"], optimized_resume["headline"]))

    original_experience = original_profile.get("work_experience") or []
    optimized_experience = optimized_resume.get("experience") or []
    for original_entry, optimized_entry in zip(original_experience, optimized_experience):
        original_bullets = original_entry.get("bullets") or []
        optimized_bullets = optimized_entry.get("bullets") or []
        for original_bullet, optimized_bullet in zip(original_bullets, optimized_bullets):
            if original_bullet and optimized_bullet:
                pairs.append((original_bullet, optimized_bullet))

    original_projects = original_profile.get("projects") or []
    optimized_projects = optimized_resume.get("projects") or []
    for original_entry, optimized_entry in zip(original_projects, optimized_projects):
        if original_entry.get("description") and optimized_entry.get("description"):
            pairs.append((original_entry["description"], optimized_entry["description"]))

    return pairs


def _best_match(paragraphs: list[Paragraph], text: str) -> int | None:
    """Index of the paragraph whose text most closely matches `text`, or None if no paragraph clears the threshold."""
    best_idx, best_ratio = None, _MATCH_THRESHOLD
    normalized = text.strip().lower()
    for i, paragraph in enumerate(paragraphs):
        ratio = SequenceMatcher(None, paragraph.text.strip().lower(), normalized).ratio()
        if ratio > best_ratio:
            best_idx, best_ratio = i, ratio
    return best_idx


def _replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    """Replace a paragraph's visible text while keeping its first run's formatting (font/size/bold/etc.)."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def render_docx(source_bytes: bytes, original_profile: dict, optimized_resume: dict) -> bytes:
    """
    Apply the LLM's tailored text onto the candidate's original uploaded
    .docx, in place, so the original fonts/spacing/layout survive untouched —
    only the matched sentences change.
    """
    document = Document(io.BytesIO(source_bytes))
    paragraphs = document.paragraphs

    for original_text, new_text in _collect_replacements(original_profile, optimized_resume):
        idx = _best_match(paragraphs, original_text)
        if idx is not None:
            _replace_paragraph_text(paragraphs[idx], new_text)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
